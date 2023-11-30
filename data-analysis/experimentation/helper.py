import docx

import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import sent_tokenize
import re

def get_paragraphs(file_path):
    # Load the document
    doc = docx.Document(file_path)

    # Iterate over each paragraph in the document
    paragraphs = []
    for para in doc.paragraphs:
        # Iif the para is empty, skip it
        paragraphs.append(para.text)
    
    return paragraphs

# Create function that loads the document and splits text into 3 sentences each. 
def get_sentences(file_path):
    # Load the document
    doc = docx.Document(file_path)

    # Iterate over each paragraph in the document
    sentences = []
    for para in doc.paragraphs:
        # Split the paragraph into sentences
        for sent in nltk.sent_tokenize(para.text):
            # Append the sentence to the list
            sentences.append(sent)
    
    return sentences


